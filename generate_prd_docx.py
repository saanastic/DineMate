from docx import Document

name = 'DineMate_Auth_Deployment_PRD.docx'
doc = Document()

# Title

doc.add_heading('DineMate Authentication and Deployment PRD', 0)

intro = doc.add_paragraph()
intro.add_run('Purpose: ').bold = True
intro.add_run(
    'Align frontend auth routing, backend API configuration, and production deployment settings to ensure DineMate login works consistently both locally and in production.'
)

doc.add_heading('Background', level=1)
doc.add_paragraph(
    'The current deployment issue is caused by the frontend bundling or environment configuration requesting localhost backend URLs in production. '
    'The login flow should target the deployed API host and use the backend auth router path /api/v1/auth/login.'
)

doc.add_heading('User Stories', level=1)
user_stories = [
    (
        'Frontend login uses deployed backend',
        'As a restaurant staff user, I want the login form to call the deployed DineMate API instead of localhost so that I can sign in from the production website.'
    ),
    (
        'Auth route works behind reverse proxy',
        'As a developer, I want the frontend to resolve /api/v1 to the production backend when deployed so that auth and data requests succeed without requiring localhost addresses.'
    ),
    (
        'CORS allows production frontend',
        'As a deployment engineer, I want the backend to accept requests from the published DineMate frontend origin so that the browser does not block login requests.'
    ),
    (
        'Shared auth page is active',
        'As a user, I want a single consistent login page that actually submits credentials to the API so I am not blocked by a stale or hidden auth screen.'
    ),
]
for title, story in user_stories:
    p = doc.add_paragraph()
    p.add_run(title + ': ').bold = True
    p.add_run(story)

doc.add_heading('Acceptance Criteria', level=1)
criteria = [
    'The frontend auth service must use VITE_API_BASE_URL or an environment-aware fallback, not a hardcoded localhost URL.',
    'The login request POST should target /api/v1/auth/login when the backend is deployed at /api/v1.',
    'Rendering should not break when the app is served from a different frontend host than the backend.',
    'CORS configuration in the backend must include the deployed frontend origin and/or allow the correct production base URL.',
    'The auth page at / or /login must be the one the router uses in production and should redirect authenticated users to /dashboard.',
    'The document should include deployment guidance for Render or equivalent with VITE_API_BASE_URL set to the production API host.',
]
for item in criteria:
    doc.add_paragraph('- ' + item, style='List Bullet')

doc.add_heading('Implementation Notes', level=1)
notes = [
    'Update client/src/services/api.js to normalize VITE_API_BASE_URL and fallback to /api/v1 when local proxy behavior is not available.',
    'Confirm auth routes in server/app/api/v1/auth.py are mounted under /api/v1/auth and that login is exposed on POST /login.',
    'Use backend environment settings (e.g. BACKEND_CORS_ORIGINS) to allow production frontend origins.',
    'Document .env or Render config values for VITE_API_BASE_URL, especially in production builds.',
]
for note in notes:
    doc.add_paragraph('- ' + note, style='List Bullet')

doc.add_heading('Deployment Verification', level=1)
verifications = [
    'Build the frontend with VITE_API_BASE_URL=https://<production-backend>/api/v1 and verify the bundle does not contain localhost URLs.',
    'Open the deployed frontend and confirm the network request for login points to the production API host and path /api/v1/auth/login.',
    'If using Render, set frontend env variable on the service and ensure backend CORS covers the frontend origin.',
]
for item in verifications:
    doc.add_paragraph('- ' + item, style='List Bullet')

doc.add_page_break()

doc.add_heading('Notes', level=1)
doc.add_paragraph(
    'If production backend and frontend are separated, ensure the frontend uses the correct runtime environment variables and the backend is reachable through the configured host instead of localhost.'
)

doc.save(name)
print(f'Saved {name}')
